from datetime import datetime
from docx.api import Document
import spacy
import pandas as pd

nlp = spacy.load('en_core_web_sm')


class NlpDataStruct:
    def __init__(self, path):
        self.path = path
        self.para_dic = {}
        self.table_dic = {}

    def run_nlp(self):
        document = Document(self.path)

        tables = document.tables
        paragraphs = document.paragraphs
        self.para_dic = {}
        counter = 0
        for para in paragraphs:
            counter += 1
            if para.text != '':
                doc = nlp(para.text)
                tokens = []
                for token in doc:
                    tokens.append({'text': token.text, 'lemma': token.lemma_, 'pos': token.pos_, 'tag': token.tag_,
                                   'dep': token.dep_, 'shape': token.shape_, 'alpha': token.is_alpha,
                                   'is_stop': token.is_stop})
                ner = []
                for ent in doc.ents:
                    ner.append([ent.text, ent.label_])
                is_ner = True if len(ner) > 0 else False
                self.para_dic[counter] = {
                    'text': para.text, 'ner': ner, 'tokens': tokens, 'is_ner': is_ner,
                    'alignment': para.alignment, 'indent_first_line': para.paragraph_format.first_line_indent,
                    'page_break': para.paragraph_format.page_break_before, 'allcaps': para.style.font.all_caps,
                    'is_bold': para.style.font.bold, 'is_italic': para.style.font.italic,
                    'color': para.style.font.color.rgb, 'underline': para.style.font.underline,
                    'fontname': para.style.font.name}
        self.table_dic = {}
        counter = 0
        for table in tables:
            data = []
            counter += 1
            keys = None
            table_frame = []
            for i, row in enumerate(table.rows):
                r = []
                for j in range(len(row.cells)):
                    r.append(row.cells[j].text)
                    if i == 0:
                        keys = tuple(row.cells[j].text)
                        continue
                    doc = nlp(row.cells[j].text)
                    tokens = []
                    for token in doc:
                        tokens.append(
                            {'text': token.text, 'lemma': token.lemma_, 'pos': token.pos_, 'tag': token.tag_,
                             'dep': token.dep_, 'shape': token.shape_, 'alpha': token.is_alpha,
                             'is_stop': token.is_stop})
                    ner = []
                    for ent in doc.ents:
                        ner.append([ent.text, ent.label_])
                    is_tab_ner = True if len(ner) > 0 else False
                    self.table_dic[counter] = {}
                    self.table_dic[counter][i] = {}
                    self.table_dic[counter][i][j] = {'row': data, 'ner': ner, 'tokens': tokens,
                                                     'is_ner': is_tab_ner}
                table_frame.append(r)
            self.table_dic[counter]['frame'] = pd.DataFrame(table_frame)
            self.table_dic[counter]['alignment'] = table.alignment
            self.table_dic[counter]['style'] = table.style
            self.table_dic[counter]['alignment'] = table.alignment


if __name__ == '__main__':
    start = datetime.now()

    path1 = './data/Sample-Rental-Agreement.docx'
    doku = NlpDataStruct(path1)
    doku.run_nlp()

    end = datetime.now()

    print('process completed in {} secs'.format(end-start))






